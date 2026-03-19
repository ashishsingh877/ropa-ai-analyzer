"""
ROPA AI Analyzer  v3.0  —  Powered by Groq (FREE)
• Whisper-large-v3  →  transcription with timestamps
• Llama-3.3-70b-versatile  →  ROPA analysis & MOM generation
• Word .docx MOM output matching Protiviti template

API key in .streamlit/secrets.toml  →  GROQ_API_KEY = "gsk_..."
"""

import streamlit as st
import groq as groq_sdk
import json, re, tempfile, os, io, time
from datetime import datetime
import pandas as pd
from mom_docx import build_mom_docx

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ROPA AI Analyzer",
    page_icon="🔐",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Styles ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@300;400;500;600;700&display=swap');
*{box-sizing:border-box;}
html,body,[data-testid="stAppViewContainer"]{background:#07090f!important;color:#cfd8ec!important;font-family:'IBM Plex Sans',sans-serif!important;}
[data-testid="stHeader"]{background:transparent!important;}
.main .block-container{max-width:1100px;padding:1.8rem 2rem 5rem;}

.app-header{padding:1.6rem 0 1.2rem;border-bottom:1px solid #161d2e;margin-bottom:1.8rem;}
.app-title{font-family:'IBM Plex Mono',monospace;font-size:1.5rem;font-weight:600;color:#eef2ff;letter-spacing:-.02em;}
.app-badge{display:inline-block;background:rgba(34,197,94,.12);color:#22c55e;font-family:'IBM Plex Mono',monospace;font-size:.62rem;font-weight:600;letter-spacing:.1em;text-transform:uppercase;padding:.18rem .65rem;border-radius:20px;border:1px solid rgba(34,197,94,.2);margin-left:.7rem;vertical-align:middle;}
.app-badge-model{display:inline-block;background:rgba(124,92,252,.12);color:#7c5cfc;font-family:'IBM Plex Mono',monospace;font-size:.62rem;font-weight:600;letter-spacing:.1em;text-transform:uppercase;padding:.18rem .65rem;border-radius:20px;border:1px solid rgba(124,92,252,.2);margin-left:.4rem;vertical-align:middle;}
.app-sub{font-size:.78rem;color:#465370;margin-top:.25rem;letter-spacing:.04em;text-transform:uppercase;}

.card{background:#0c0f1a;border:1px solid #161d2e;border-radius:12px;padding:1.3rem 1.4rem;margin-bottom:1.3rem;}
.card-label{font-family:'IBM Plex Mono',monospace;font-size:.63rem;color:#4f8ef7;letter-spacing:.12em;text-transform:uppercase;margin-bottom:.4rem;display:flex;align-items:center;gap:.45rem;}
.card-label::before{content:'';display:inline-block;width:5px;height:5px;border-radius:50%;background:#4f8ef7;}
.card-title{font-size:.92rem;font-weight:600;color:#b8cae0;margin-bottom:.8rem;}

.compat-row{display:flex;gap:.5rem;flex-wrap:wrap;margin-bottom:.8rem;}
.compat-pill{display:inline-flex;align-items:center;gap:.35rem;padding:.25rem .75rem;border-radius:20px;font-size:.72rem;font-weight:600;letter-spacing:.04em;}
.pill-green{background:rgba(34,197,94,.1);color:#22c55e;border:1px solid rgba(34,197,94,.2);}
.pill-blue{background:rgba(79,142,247,.1);color:#4f8ef7;border:1px solid rgba(79,142,247,.2);}
.pill-purple{background:rgba(124,92,252,.1);color:#7c5cfc;border:1px solid rgba(124,92,252,.2);}
.pill-amber{background:rgba(245,158,11,.1);color:#f59e0b;border:1px solid rgba(245,158,11,.2);}

.stButton>button{background:linear-gradient(135deg,#22c55e,#4f8ef7)!important;color:#fff!important;border:none!important;border-radius:8px!important;font-family:'IBM Plex Sans',sans-serif!important;font-weight:600!important;font-size:.86rem!important;padding:.5rem 1.5rem!important;transition:opacity .2s!important;}
.stButton>button:hover{opacity:.84!important;}
.stButton>button:disabled{opacity:.32!important;}

.stTextArea textarea,.stTextInput input{background:#0b0e19!important;border:1px solid #161d2e!important;border-radius:8px!important;color:#c0cce0!important;font-family:'IBM Plex Mono',monospace!important;font-size:.79rem!important;}
.stTextArea textarea:focus,.stTextInput input:focus{border-color:#22c55e!important;box-shadow:0 0 0 2px rgba(34,197,94,.12)!important;}
[data-testid="stFileUploader"]{background:#0b0e19!important;border:1.5px dashed #1e2840!important;border-radius:10px!important;padding:.7rem!important;}
[data-testid="stFileUploader"]:hover{border-color:#22c55e!important;}
.stSelectbox>div>div{background:#0b0e19!important;border:1px solid #161d2e!important;border-radius:8px!important;color:#c0cce0!important;}

.stTabs [data-baseweb="tab-list"]{gap:.3rem;background:transparent!important;border-bottom:1px solid #161d2e!important;}
.stTabs [data-baseweb="tab"]{background:transparent!important;color:#465370!important;border-radius:6px 6px 0 0!important;font-family:'IBM Plex Sans',sans-serif!important;font-size:.82rem!important;padding:.42rem .85rem!important;}
.stTabs [aria-selected="true"]{background:#0c0f1a!important;color:#22c55e!important;border:1px solid #161d2e!important;border-bottom:1px solid #0c0f1a!important;}

hr{border-color:#161d2e!important;margin:1.6rem 0!important;}
[data-testid="stExpander"]{background:#0c0f1a!important;border:1px solid #161d2e!important;border-radius:8px!important;}

.qa-card{background:#0b0e19;border:1px solid #161d2e;border-left:3px solid #4f8ef7;border-radius:8px;padding:1rem 1.2rem;margin-bottom:.85rem;}
.qa-card.c-hi{border-left-color:#22c55e;}.qa-card.c-md{border-left-color:#f59e0b;}.qa-card.c-lo{border-left-color:#ef4444;}
.qa-num{font-family:'IBM Plex Mono',monospace;font-size:.65rem;color:#465370;text-transform:uppercase;letter-spacing:.08em;margin-bottom:.3rem;}
.qa-q{font-size:.83rem;font-weight:600;color:#7a92b8;margin-bottom:.45rem;}
.qa-a{font-size:.88rem;color:#bfcde0;line-height:1.6;margin-bottom:.6rem;}
.qa-meta{display:flex;gap:1rem;flex-wrap:wrap;font-family:'IBM Plex Mono',monospace;font-size:.68rem;color:#465370;}
.badge{display:inline-flex;align-items:center;padding:.13rem .5rem;border-radius:20px;font-size:.62rem;font-weight:600;letter-spacing:.06em;text-transform:uppercase;}
.b-hi{background:rgba(34,197,94,.1);color:#22c55e;}.b-md{background:rgba(245,158,11,.1);color:#f59e0b;}.b-lo{background:rgba(239,68,68,.1);color:#ef4444;}.b-t{background:rgba(79,142,247,.1);color:#4f8ef7;}
.verbatim{background:#070912;border-left:2px solid #1a2235;border-radius:4px;padding:.45rem .75rem;margin-top:.45rem;font-family:'IBM Plex Mono',monospace;font-size:.74rem;color:#5a7090;line-height:1.5;font-style:italic;}

.stat-grid{display:grid;grid-template-columns:repeat(5,1fr);gap:.7rem;margin-bottom:1.3rem;}
.stat-box{background:#0c0f1a;border:1px solid #161d2e;border-radius:10px;padding:.9rem;text-align:center;}
.sn{font-family:'IBM Plex Mono',monospace;font-size:1.5rem;font-weight:600;color:#22c55e;}
.sl{font-size:.7rem;color:#465370;margin-top:.15rem;}

.ok{background:rgba(34,197,94,.07);border:1px solid rgba(34,197,94,.2);border-radius:8px;padding:.65rem 1rem;color:#22c55e;font-size:.83rem;margin-bottom:.9rem;}
.warn{background:rgba(245,158,11,.07);border:1px solid rgba(245,158,11,.2);border-radius:8px;padding:.65rem 1rem;color:#f59e0b;font-size:.83rem;margin-bottom:.9rem;}
.info{background:rgba(79,142,247,.07);border:1px solid rgba(79,142,247,.2);border-radius:8px;padding:.65rem 1rem;color:#4f8ef7;font-size:.83rem;margin-bottom:.9rem;}
.sh{font-family:'IBM Plex Mono',monospace!important;font-size:.7rem!important;color:#22c55e!important;letter-spacing:.12em!important;text-transform:uppercase!important;padding-bottom:.45rem!important;border-bottom:1px solid #161d2e!important;margin:1.6rem 0 .9rem!important;}

.mom-box{background:#0b0e19;border:1px solid #161d2e;border-radius:8px;padding:.9rem 1.2rem;margin-bottom:.75rem;}
.mom-lbl{font-family:'IBM Plex Mono',monospace;font-size:.67rem;color:#7c5cfc;letter-spacing:.1em;text-transform:uppercase;margin-bottom:.5rem;}
.mom-val{font-size:.87rem;color:#adbbd4;line-height:1.65;}
.ai-card{background:#0b0e19;border:1px solid #161d2e;border-left:3px solid #7c5cfc;border-radius:8px;padding:.9rem 1.1rem;margin-bottom:.75rem;}

/* groq speed indicator */
.groq-banner{background:linear-gradient(135deg,rgba(34,197,94,.08),rgba(79,142,247,.08));border:1px solid rgba(34,197,94,.2);border-radius:10px;padding:.9rem 1.2rem;margin-bottom:1.3rem;display:flex;align-items:center;gap:1rem;flex-wrap:wrap;}
.groq-icon{font-size:1.4rem;}
.groq-text{font-size:.83rem;color:#adbbd4;line-height:1.5;}
.groq-text b{color:#22c55e;}
</style>
""", unsafe_allow_html=True)

# ── Helpers ────────────────────────────────────────────────────────────────────
def fmt_time(s):
    if s is None: return "N/A"
    s=int(s); h,r=divmod(s,3600); m,sec=divmod(r,60)
    return f"{h:02d}:{m:02d}:{sec:02d}" if h else f"{m:02d}:{sec:02d}"

def get_client():
    try:
        key = st.secrets["GROQ_API_KEY"]
        return groq_sdk.Groq(api_key=key)
    except KeyError:
        st.error("⛔ **GROQ_API_KEY** not found in Streamlit secrets.\n\n"
                 "Get your free key at **console.groq.com** → API Keys → Create key\n\n"
                 "Then add it in `.streamlit/secrets.toml` as:\n```\nGROQ_API_KEY = \"gsk_...\"\n```")
        st.stop()

# ════════════════════════════════════════════════════════════════════════════
# ROPA TEMPLATE PARSER  —  handles complex multi-sheet Excel ROPA templates
# Strategy: use openpyxl directly for proper merged-cell support,
#           walk every section, collect EVERY field label across ALL columns.
# ════════════════════════════════════════════════════════════════════════════

# Cell values that are structural noise, not ROPA fields
_SKIP_EXACT = {
    'nan','none','','yes','no','n/a','na','tbd','-','—',
    'remarks','comments','note','notes','ref','reference',
    'true','false','s.no','sr.no','sr no','serial no',
}

def _is_field_id(text: str) -> bool:
    """True for structural field-id codes: 1.1  /  2.3  /  1.1.2"""
    return bool(re.match(r'^\d+(\.\d+){0,3}\.?$', text.strip()))

def _is_section_header(text: str) -> bool:
    """True for rows like 'Section 1: Process Overview'"""
    return bool(re.match(r'section\s*\d', text.strip(), re.I))

def _clean_cell(v) -> str:
    """Normalise a cell value to a clean string."""
    if v is None: return ''
    t = str(v).strip()
    # Remove common Excel artefacts
    t = re.sub(r'[\r\n]+', ' ', t)   # newlines → space
    t = re.sub(r' {2,}', ' ', t)       # multiple spaces → one
    return t.strip()

def _is_good_field(text: str) -> bool:
    """Return True if this cell looks like a genuine ROPA field label."""
    t = text.strip()
    if not t or len(t) < 3: return False
    if t.lower() in _SKIP_EXACT: return False
    if _is_field_id(t): return False
    if _is_section_header(t): return False
    # Pure numeric
    if re.match(r'^[\d,.]+$', t): return False
    # Very generic single words that are not fields
    if t.lower() in {'country','date','name','no','yes','item',
                     'description','details','value','type','code'}:
        # Allow them only if they have context (section prefix added below)
        pass
    # Truncate very long instructional text but still KEEP it —
    # many ROPA fields have long descriptions like
    # "Specify the type of data subject (Customers, Vendor, employee etc.)"
    return True


def _read_xlsx_with_openpyxl(raw_bytes: bytes) -> list:
    """
    Read ALL ROPA field labels from an Excel file using openpyxl.

    This handles:
    • Merged cells  (field labels often span multiple rows)
    • Multiple sheets  (skip glossary/lookup/cover sheets)
    • Multi-section layout  (Section 1, Section 2 … N)
    • Field labels spread horizontally across columns
    • Long descriptive cells (keep, truncate at 250 chars)

    Returns a list of strings like:
      "[Section 1: Process Overview] Country"
      "[Section 1: Process Overview] Function Name"
      "[Section 2: Data Elements] Type of Personal Data Collected"
    """
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=True)

    SKIP_SHEET_PAT = re.compile(
        r'(glossary|lookup|ref|list|dropdown|master|cover|index|'
        r'instruction|guide|readme|changelog|version)', re.I
    )

    all_questions = []
    seen_keys     = set()

    for sheet_name in wb.sheetnames:
        if SKIP_SHEET_PAT.search(sheet_name):
            continue

        ws = wb[sheet_name]

        # Build a 2-D grid of merged-cell-aware values
        # openpyxl fills merged cells with None for non-anchors;
        # we propagate the anchor value across the merge range.
        grid = {}  # (row, col) → string value

        # First pass: real cell values
        for row in ws.iter_rows():
            for cell in row:
                v = _clean_cell(cell.value)
                grid[(cell.row, cell.column)] = v

        # Second pass: fill merged ranges with anchor value
        for merge in ws.merged_cells.ranges:
            anchor_val = grid.get((merge.min_row, merge.min_col), '')
            for r in range(merge.min_row, merge.max_row + 1):
                for c in range(merge.min_col, merge.max_col + 1):
                    if not grid.get((r, c)):
                        grid[(r, c)] = anchor_val

        max_row = ws.max_row or 1
        max_col = ws.max_column or 1

        current_section = sheet_name   # default context = sheet name

        row_num = 1
        while row_num <= max_row:
            # Collect all non-empty cells in this row
            row_cells = []
            for c in range(1, max_col + 1):
                v = grid.get((row_num, c), '')
                if v:
                    row_cells.append(v)

            if not row_cells:
                row_num += 1
                continue

            # ── Detect section header row ─────────────────────────────────
            # A row is a section header if ≥1 cell matches "Section N…"
            # and most other cells are empty or field-IDs
            section_cells = [c for c in row_cells if _is_section_header(c)]
            if section_cells:
                current_section = section_cells[0]
                row_num += 1
                continue

            # ── Detect field-ID row ───────────────────────────────────────
            # Rows like:  1.1 | 1.2 | 1.3 | 1.4  (all cells are field IDs)
            field_id_cells = [c for c in row_cells if _is_field_id(c)]
            if field_id_cells and len(field_id_cells) >= len(row_cells) * 0.5:
                # This is a header/numbering row — skip it
                row_num += 1
                continue

            # ── Collect field labels from this row ────────────────────────
            for val in row_cells:
                if not _is_good_field(val):
                    continue
                # Truncate very long values but keep first 200 chars
                label_text = val[:200].strip()
                if not label_text:
                    continue
                # Build the full question: [Section context] Field label
                question = f"[{current_section}] {label_text}"
                key = label_text.lower()[:100]  # dedup key (first 100 chars)
                if key not in seen_keys:
                    seen_keys.add(key)
                    all_questions.append(question)

            row_num += 1

    return all_questions


def parse_questions(text: str) -> list:
    """Extract questions from free-form text (numbered, bulleted, or plain)."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    pat   = re.compile(r'^(?:Q(?:uestion)?\s*)?(\d+)[.):–\-]\s*(.+)', re.I)
    qs    = []
    for line in lines:
        m = pat.match(line)
        if m:   qs.append(m.group(2).strip())
        elif line.startswith(('-','•','*','–')): qs.append(line.lstrip('-•*– ').strip())
        else:   qs.append(line)
    seen, out = set(), []
    for q in qs:
        q = re.sub(r'[ \t]+', ' ', q).strip()
        if q and len(q) >= 3 and q.lower() not in seen:
            seen.add(q.lower()); out.append(q)
    return out


def parse_template_file(f) -> list:
    """Parse ALL ROPA field labels from an uploaded file (all sheets, all sections)."""
    name = f.name.lower()
    raw  = f.read()

    # ── TXT ──────────────────────────────────────────────────────────────────
    if name.endswith('.txt'):
        return parse_questions(raw.decode('utf-8', 'ignore'))

    # ── CSV ──────────────────────────────────────────────────────────────────
    if name.endswith('.csv'):
        try:
            df = pd.read_csv(io.BytesIO(raw), header=None, dtype=str)
            qs = []
            seen_k = set()
            for _, row in df.iterrows():
                for v in row:
                    t = re.sub(r' +',' ', str(v).strip()) if pd.notna(v) else ''
                    if _is_good_field(t) and t.lower() not in seen_k:
                        seen_k.add(t.lower()); qs.append(t)
            return qs if qs else parse_questions(df.to_csv(index=False))
        except Exception as e:
            st.error(f"CSV parse error: {e}"); return []

    # ── XLSX / XLS ───────────────────────────────────────────────────────────
    if name.endswith(('.xlsx', '.xls')):
        try:
            qs = _read_xlsx_with_openpyxl(raw)
            if not qs:
                # Hard fallback: pandas flat scan
                xl = pd.ExcelFile(io.BytesIO(raw))
                seen_k = set()
                for sheet in xl.sheet_names:
                    df = pd.read_excel(xl, sheet_name=sheet, header=None, dtype=str)
                    for _, row in df.iterrows():
                        for v in row:
                            t = str(v).strip() if pd.notna(v) else ''
                            if _is_good_field(t) and t.lower() not in seen_k:
                                seen_k.add(t.lower()); qs.append(t)
            return qs
        except Exception as e:
            st.error(f"Excel parse error: {e}"); return []

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if name.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            seen_k, out = set(), []
            def _add(t):
                t = re.sub(r' +',' ',t.strip())
                k = t.lower()
                if t and _is_good_field(t) and k not in seen_k:
                    seen_k.add(k); out.append(t)
            for p in doc.paragraphs: _add(p.text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells: _add(cell.text)
            return out
        except Exception as e:
            st.error(f"DOCX parse error: {e}"); return []

    # ── PDF ──────────────────────────────────────────────────────────────────
    if name.endswith('.pdf'):
        try:
            from pypdf import PdfReader
            rdr  = PdfReader(io.BytesIO(raw))
            text = "".join(p.extract_text() or "" for p in rdr.pages)
            return parse_questions(text)
        except Exception as e:
            st.error(f"PDF parse error: {e}"); return []

    st.warning(f"Unsupported file type: {name}")
    return []


# ── Transcription (Groq Whisper-large-v3) ────────────────────────────────────

GROQ_MAX_BYTES = 24 * 1024 * 1024   # 24 MB (Groq hard limit is 25 MB)

# MIME type map for correct Content-Type headers
MIME_MAP = {
    ".mp3":  "audio/mpeg",
    ".mp4":  "video/mp4",
    ".m4a":  "audio/mp4",
    ".mpeg": "audio/mpeg",
    ".mpga": "audio/mpeg",
    ".wav":  "audio/wav",
    ".webm": "audio/webm",
    ".ogg":  "audio/ogg",
    ".flac": "audio/flac",
}


def _normalise_resp(resp) -> dict:
    """Convert Groq response object → plain dict with text + segments list."""
    if hasattr(resp, 'model_dump'):
        data = resp.model_dump()
    else:
        data = {
            "text":     str(getattr(resp, 'text', resp)),
            "segments": list(getattr(resp, 'segments', []) or []),
        }
    raw_segs = data.get("segments") or []
    segs = []
    for s in raw_segs:
        if isinstance(s, dict):
            segs.append(s)
        elif hasattr(s, 'model_dump'):
            segs.append(s.model_dump())
        else:
            segs.append({"text": str(s), "start": None, "end": None})
    data["segments"] = segs
    return data


def _send_to_groq(path: str, client, suffix: str) -> dict:
    """Send a single audio file path to Groq Whisper. Returns normalised dict."""
    mime = MIME_MAP.get(suffix, "audio/mpeg")
    with open(path, "rb") as f:
        resp = client.audio.transcriptions.create(
            model="whisper-large-v3",
            file=(os.path.basename(path), f, mime),
            response_format="verbose_json",
            timestamp_granularities=["segment"],
            temperature=0.0,          # deterministic = more accurate
        )
    return _normalise_resp(resp)


def _get_duration_ffprobe(src_path: str) -> float:
    """Return audio duration in seconds using ffprobe."""
    import subprocess
    cmd = ["ffprobe","-v","error","-show_entries","format=duration",
           "-of","default=noprint_wrappers=1:nokey=1", src_path]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
    try:
        return float(result.stdout.strip())
    except (ValueError, AttributeError):
        return 3600.0   # fallback: assume 60 min


def _split_audio_ffmpeg(src_path: str, chunk_sec: int = 600) -> list:
    """
    Split audio into chunks using ffmpeg subprocess directly.
    No pydub / pyaudioop needed — works on Python 3.13+.
    Returns list of (tmp_mp3_path, start_offset_seconds).
    """
    import subprocess
    duration = _get_duration_ffprobe(src_path)
    chunks, offset = [], 0.0
    while offset < duration:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        tmp.close()
        cmd = ["ffmpeg","-y","-ss",str(offset),"-t",str(chunk_sec),
               "-i",src_path,"-ar","16000","-ac","1","-b:a","64k",
               "-f","mp3", tmp.name]
        res = subprocess.run(cmd, capture_output=True, timeout=180)
        if res.returncode != 0:
            raise RuntimeError(res.stderr.decode("utf-8",errors="ignore")[-400:])
        chunks.append((tmp.name, offset))
        offset += chunk_sec
    return chunks


def transcribe_audio(audio_bytes: bytes, filename: str, client) -> dict:
    """
    Transcribe audio using Groq Whisper-large-v3.
    • Files ≤ 24 MB  → sent directly (single request).
    • Files  > 24 MB → split into 10-minute MP3 chunks via ffmpeg,
                       each chunk transcribed separately, then stitched.
    Correct approach: ALWAYS split at proper audio boundaries, never raw bytes.
    """
    suffix = os.path.splitext(filename)[-1].lower() or ".mp3"
    if suffix not in MIME_MAP:
        suffix = ".mp3"

    # Write full file to temp path (needed for ffmpeg and direct send)
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(audio_bytes)
        src_path = tmp.name

    try:
        file_size = len(audio_bytes)

        # ── Small file: send directly ─────────────────────────────────────
        if file_size <= GROQ_MAX_BYTES:
            return _send_to_groq(src_path, client, suffix)

        # ── Large file: split properly with ffmpeg ─────────────────────────
        st.markdown(
            '<div class="info">📦 File is large — splitting into 10-min chunks '
            'via ffmpeg (this is normal for Teams/Meet recordings).</div>',
            unsafe_allow_html=True
        )

        try:
            chunks = _split_audio_ffmpeg(src_path, chunk_sec=600)
        except Exception as split_err:
            st.error(
                f"❌ Audio splitting failed: {split_err}\n\n"
                "**Fix:** ffmpeg must be installed on the server.\n"
                "- **Local:** `brew install ffmpeg` (Mac) or `sudo apt install ffmpeg` (Linux)\n"
                "- **Streamlit Cloud:** make sure `packages.txt` contains `ffmpeg`"
            )
            st.stop()

        full_text   = ""
        all_segs    = []

        for i, (chunk_path, time_offset) in enumerate(chunks):
            chunk_size = os.path.getsize(chunk_path)
            status_msg = f"Transcribing chunk {i+1}/{len(chunks)}…"
            st.caption(f"  ⏳ {status_msg}")
            try:
                d = _send_to_groq(chunk_path, client, ".mp3")
            except Exception as chunk_err:
                st.warning(f"⚠️ Chunk {i+1} failed ({chunk_err}) — skipping.")
                d = {"text": "", "segments": []}
            finally:
                try: os.unlink(chunk_path)
                except: pass

            full_text += (" " if full_text else "") + (d.get("text") or "")

            for s in d.get("segments", []):
                sd = dict(s)
                if sd.get("start") is not None:
                    sd["start"] = round(sd["start"] + time_offset, 2)
                if sd.get("end") is not None:
                    sd["end"]   = round(sd["end"]   + time_offset, 2)
                all_segs.append(sd)

        return {"text": full_text.strip(), "segments": all_segs}

    finally:
        try: os.unlink(src_path)
        except: pass

# ── ROPA Analysis (Groq Llama-3.3-70b) ───────────────────────────────────────
def analyze_ropa(transcript: str, segments: list, questions: list, client) -> dict:
    seg_block = ""
    if segments:
        seg_block = "\n\nTIMESTAMPED SEGMENTS (use for accurate timestamps):\n"
        for s in segments[:150]:
            seg_block += f"[{fmt_time(s.get('start'))} → {fmt_time(s.get('end'))}] {s.get('text','').strip()}\n"

    system = """You are a senior Data Protection Officer AI. Extract precise ROPA answers from a meeting transcript.

RULES:
- Be SPECIFIC: extract exact system names, legal bases, retention periods, data categories — never generalise.
- For each question, find the answer in the transcript and give the timestamp range (MM:SS).
- Confidence: HIGH = explicit clear answer | MEDIUM = inferred or partial | LOW = not discussed.
- Include a verbatim quote (max 40 words) from the transcript as evidence for the human reviewer.
- If a question was NOT answered, say so clearly with confidence=LOW.

You MUST return ONLY valid JSON, no markdown, no explanation, matching this exact schema:
{
  "answers": [
    {
      "question_index": 0,
      "question": "...",
      "answer": "Detailed accurate answer from transcript",
      "confidence": "HIGH",
      "timestamp_start": "MM:SS",
      "timestamp_end": "MM:SS",
      "verbatim_quote": "Exact words from transcript, max 40 words",
      "notes": "Any caveats or follow-up needed"
    }
  ],
  "overall_completeness": 80,
  "summary": "One paragraph summary of ROPA coverage"
}"""

    q_block = "\n".join(f"{i+1}. {q}" for i,q in enumerate(questions))

    # Groq context window is 128k tokens for llama-3.3-70b — plenty for most meetings
    # Truncate transcript to ~80k chars to stay safe
    transcript_trunc = transcript[:80000]
    user = f"ROPA QUESTIONS:\n{q_block}\n\nFULL TRANSCRIPT:\n{transcript_trunc}{seg_block}"

    resp = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role":"system","content":system},
            {"role":"user","content":user}
        ],
        temperature=0.1,
        max_tokens=4096,
        response_format={"type":"json_object"}
    )
    raw = resp.choices[0].message.content
    # Strip any accidental markdown fences
    raw = re.sub(r'^```(?:json)?\s*','',raw.strip()); raw = re.sub(r'```\s*$','',raw)
    return json.loads(raw)

# ── MOM Generation (Groq Llama-3.3-70b) ──────────────────────────────────────
def generate_mom(transcript: str, segments: list, meta: dict, client) -> dict:
    seg_block = ""
    if segments:
        seg_block = "\n\nKEY TIMESTAMPS:\n" + "\n".join(
            f"[{fmt_time(s.get('start'))}] {s.get('text','').strip()}" for s in segments[:100]
        )

    system = """You are a professional meeting secretary. Generate thorough, structured Meeting Minutes.
Extract every decision, action item, and unresolved question. Include timestamps where available.

Return ONLY valid JSON (no markdown fences):
{
  "meeting_title": "...",
  "agenda": "Brief one-line agenda",
  "attendees_mentioned": ["Name or Role from transcript"],
  "key_discussion_points": [
    {"point": "Short topic title", "detail": "Full discussion detail", "timestamp": "MM:SS or null"}
  ],
  "decisions_made": ["Exact decision 1"],
  "action_items": [
    {"action": "What must be done", "owner": "Person or TBD", "due_date": "Date or TBD", "timestamp": "MM:SS or null"}
  ],
  "questions_raised": ["Unresolved question 1"],
  "next_steps": "Paragraph summarising what happens next",
  "next_meeting": "Date/time if mentioned, else null"
}"""

    transcript_trunc = transcript[:60000]
    user = f"Meeting context: {json.dumps(meta)}\n\nTRANSCRIPT:\n{transcript_trunc}{seg_block}"

    resp = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role":"system","content":system},
            {"role":"user","content":user}
        ],
        temperature=0.15,
        max_tokens=3000,
        response_format={"type":"json_object"}
    )
    raw = resp.choices[0].message.content
    raw = re.sub(r'^```(?:json)?\s*','',raw.strip()); raw = re.sub(r'```\s*$','',raw)
    return json.loads(raw)

# ── HTML Verification Report ───────────────────────────────────────────────────
def build_verification_html(result, questions, meta):
    answers = result.get("answers",[])
    compl   = result.get("overall_completeness",0)
    summary = result.get("summary","")
    high = sum(1 for a in answers if a.get("confidence")=="HIGH")
    med  = sum(1 for a in answers if a.get("confidence")=="MEDIUM")
    low  = sum(1 for a in answers if a.get("confidence")=="LOW")

    rows=""
    for a in answers:
        conf = a.get("confidence","LOW")
        bg = {"HIGH":"#052e16","MEDIUM":"#2d1f03","LOW":"#2d0707"}.get(conf,"#111")
        cl = {"HIGH":"#22c55e","MEDIUM":"#f59e0b","LOW":"#ef4444"}.get(conf,"#aaa")
        ts = f"{a.get('timestamp_start','—')} → {a.get('timestamp_end','—')}"
        vq = str(a.get('verbatim_quote','')).replace('<','&lt;').replace('>','&gt;')
        vq_html = (f'<div style="margin-top:.5rem;background:#070912;border-left:2px solid #1a2235;'
                   f'padding:.4rem .7rem;font-family:monospace;font-size:.73rem;color:#5a7090;'
                   f'font-style:italic;">" {vq} "</div>') if vq else ''
        notes_html = (f'<div style="margin-top:.3rem;font-size:.75rem;color:#465370;">📎 {a.get("notes","")}</div>'
                      if a.get("notes") else '')
        rows += f"""<tr>
          <td style="color:#22c55e;font-family:monospace;font-size:.73rem;padding:.7rem .9rem;vertical-align:top;border-bottom:1px solid #161d2e;">Q{a.get('question_index',0)+1}</td>
          <td style="padding:.7rem .9rem;vertical-align:top;border-bottom:1px solid #161d2e;font-size:.8rem;color:#7a92b8;">{a.get('question','')}</td>
          <td style="padding:.7rem .9rem;vertical-align:top;border-bottom:1px solid #161d2e;font-size:.83rem;color:#bfcde0;">{a.get('answer','')}{vq_html}{notes_html}</td>
          <td style="padding:.7rem .9rem;vertical-align:top;border-bottom:1px solid #161d2e;font-family:monospace;font-size:.72rem;color:#4f8ef7;white-space:nowrap;">{ts}</td>
          <td style="padding:.7rem .9rem;vertical-align:top;border-bottom:1px solid #161d2e;text-align:center;">
            <span style="background:{bg};color:{cl};padding:.2rem .55rem;border-radius:20px;font-size:.62rem;font-weight:700;letter-spacing:.06em;text-transform:uppercase;">{conf}</span>
          </td>
        </tr>"""

    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>ROPA Verification Report</title>
<style>
body{{font-family:'Segoe UI',sans-serif;background:#07090f;color:#cfd8ec;margin:0;padding:2rem;}}
h1{{font-family:monospace;font-size:1.3rem;color:#eef2ff;margin-bottom:.2rem;}}
.sub{{font-size:.75rem;color:#465370;text-transform:uppercase;letter-spacing:.06em;margin-bottom:1.6rem;}}
.model-tag{{display:inline-block;background:rgba(34,197,94,.1);color:#22c55e;font-family:monospace;font-size:.65rem;padding:.15rem .5rem;border-radius:12px;margin-left:.5rem;}}
.stats{{display:flex;gap:.8rem;margin-bottom:1.6rem;flex-wrap:wrap;}}
.s{{background:#0c0f1a;border:1px solid #161d2e;border-radius:10px;padding:.8rem 1.2rem;text-align:center;min-width:105px;}}
.sn{{font-family:monospace;font-size:1.35rem;font-weight:700;color:#22c55e;}}
.sl{{font-size:.68rem;color:#465370;}}
.sumbox{{background:#0c0f1a;border:1px solid #161d2e;border-left:3px solid #22c55e;border-radius:8px;padding:.85rem 1.1rem;margin-bottom:1.4rem;font-size:.85rem;color:#adbbd4;line-height:1.6;}}
table{{width:100%;border-collapse:collapse;background:#0b0e19;border:1px solid #161d2e;border-radius:8px;overflow:hidden;font-size:.84rem;}}
thead tr{{background:#0c0f1a;}}
th{{padding:.6rem .9rem;text-align:left;font-size:.67rem;color:#22c55e;font-family:monospace;text-transform:uppercase;letter-spacing:.08em;border-bottom:1px solid #1e2840;}}
.foot{{margin-top:2rem;font-size:.7rem;color:#1e2840;text-align:center;}}
</style></head><body>
<h1>🔐 ROPA Verification Report <span class="model-tag">Groq · Llama 3.3 70B + Whisper-large-v3</span></h1>
<div class="sub">Generated {datetime.now().strftime('%d %b %Y · %H:%M')} &nbsp;·&nbsp; {meta.get('filename','recording')}</div>
<div class="sumbox">{summary}</div>
<div class="stats">
  <div class="s"><div class="sn">{len(questions)}</div><div class="sl">Questions</div></div>
  <div class="s"><div class="sn" style="color:#22c55e">{high}</div><div class="sl">High Conf.</div></div>
  <div class="s"><div class="sn" style="color:#f59e0b">{med}</div><div class="sl">Medium</div></div>
  <div class="s"><div class="sn" style="color:#ef4444">{low}</div><div class="sl">Low/Missing</div></div>
  <div class="s"><div class="sn">{compl}%</div><div class="sl">Complete</div></div>
</div>
<table><thead><tr><th>#</th><th>Question</th><th>AI Answer + Evidence</th><th>Timestamp</th><th>Confidence</th></tr></thead>
<tbody>{rows}</tbody></table>
<div class="foot">ROPA AI Analyzer v3 · Groq Free API · For internal / audit use only</div>
</body></html>"""

# ── Session state init ────────────────────────────────────────────────────────
for k,v in {"transcript":None,"segments":[],"questions":[],"ropa_result":None,
            "mom_result":None,"audio_meta":{}}.items():
    if k not in st.session_state: st.session_state[k]=v

# ── Header ─────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-header">
  <div class="app-title">ROPA AI Analyzer
    <span class="app-badge">✦ FREE — Groq API</span>
    <span class="app-badge-model">Llama 3.3 · Whisper v3</span>
  </div>
  <div class="app-sub">Privacy compliance · ROPA auto-fill · Word Meeting Minutes · Zero API cost</div>
</div>
""", unsafe_allow_html=True)

# Groq benefits banner
st.markdown("""
<div class="groq-banner">
  <div class="groq-icon">⚡</div>
  <div class="groq-text">
    <b>Running on Groq — completely free.</b>
    Whisper-large-v3 for transcription · Llama 3.3 70B for ROPA analysis & MOM generation.
    Groq's LPU hardware is <b>10–20× faster</b> than standard GPU inference.
    Get your free API key at <b>console.groq.com</b> → no credit card needed.
  </div>
</div>
""", unsafe_allow_html=True)

# Recording sources
st.markdown("""
<div class="card">
<div class="card-label">Supported Recording Sources</div>
<div class="compat-row">
  <span class="compat-pill pill-green">✓ Microsoft Teams</span>
  <span class="compat-pill pill-blue">✓ Google Meet</span>
  <span class="compat-pill pill-purple">✓ Zoom</span>
  <span class="compat-pill pill-amber">✓ In-Person / Phone</span>
  <span class="compat-pill pill-green">✓ Webex / Any MP3/MP4</span>
</div>
<div style="font-size:.78rem;color:#465370;line-height:1.65;">
  <b style="color:#7a92b8;">Teams:</b> Chat / Calendar → recording thumbnail → ··· → <b>Download (.mp4)</b><br>
  <b style="color:#7a92b8;">Google Meet:</b> Recording goes to Google Drive → right-click → <b>Download (.mp4)</b><br>
  <b style="color:#7a92b8;">Zoom:</b> Recordings tab → local .mp4/.m4a file, or cloud download from zoom.us<br>
  <b style="color:#7a92b8;">Max size:</b> Files up to 200 MB. Files over 24 MB are auto-split into chunks.
</div>
</div>
""", unsafe_allow_html=True)

# ── Tabs ───────────────────────────────────────────────────────────────────────
tab1, tab2 = st.tabs(["🔐  ROPA Analyzer", "📋  Meeting Minutes (.docx)"])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — ROPA
# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    c1,c2 = st.columns(2, gap="large")

    with c1:
        st.markdown('<div class="card"><div class="card-label">Step 1 · Recording</div>'
                    '<div class="card-title">Upload Meeting / Interview Audio</div>', unsafe_allow_html=True)
        af = st.file_uploader("MP3 · MP4 · M4A · WAV · OGG · WEBM · FLAC",
                              type=["mp3","mp4","m4a","wav","ogg","webm","flac","mpeg","mpga"],
                              key="af_ropa")
        if af:
            st.audio(af)
            st.session_state.audio_meta = {"filename":af.name,"size_mb":round(af.size/1e6,2)}
            size_warn = " ⚠️ Will be auto-chunked" if af.size > 24*1024*1024 else ""
            st.markdown(f'<div class="ok">✓ {af.name} ({st.session_state.audio_meta["size_mb"]} MB){size_warn}</div>',
                        unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="card"><div class="card-label">Step 2 · ROPA Template</div>'
                    '<div class="card-title">Upload Template or Paste Questions</div>', unsafe_allow_html=True)
        tf = st.file_uploader("TXT · CSV · XLSX · DOCX · PDF",
                              type=["txt","csv","xlsx","xls","docx","pdf"], key="tf_ropa")
        if tf:
            with st.spinner(f"Scanning all sheets & sections in {tf.name}…"):
                qs = parse_template_file(tf)
            if qs:
                st.session_state.questions = qs
                # Count how many distinct sections were found
                sections = set()
                for q in qs:
                    if q.startswith('[') and ']' in q:
                        sections.add(q.split(']')[0].lstrip('['))
                sec_info = f" across {len(sections)} sections" if sections else ""
                st.markdown(
                    f'<div class="ok">✓ Found <b>{len(qs)} field labels</b>{sec_info} in {tf.name}</div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    f'<div class="warn">⚠️ Could not extract fields from <b>{tf.name}</b>. '
                    f'Try pasting questions manually below.</div>',
                    unsafe_allow_html=True
                )
                try:
                    if tf.name.lower().endswith(('.xlsx','.xls')):
                        tf.seek(0); raw2 = tf.read()
                        _df = pd.read_excel(io.BytesIO(raw2), header=None, dtype=str, nrows=8)
                        st.caption("Preview of your file (first 8 rows):")
                        st.dataframe(_df, use_container_width=True)
                except Exception:
                    pass

        manual = st.text_area(
            "Or paste / type questions here (one per line):",
            height=160,
            placeholder="1. What personal data is processed?\n2. What is the legal basis for processing?\n3. What is the retention period?\n4. Who are the data subjects?\n5. Are there any third-party data transfers?"
        )
        if manual.strip():
            mqs = parse_questions(manual)
            if mqs:
                st.session_state.questions = mqs
                st.markdown(f'<div class="ok">✓ {len(mqs)} questions ready from text input</div>',
                            unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    # ── Status checklist + preview ─────────────────────────────────────────
    has_audio     = af is not None
    has_questions = len(st.session_state.questions) > 0

    if has_questions:
        with st.expander(f"📋 Preview {len(st.session_state.questions)} questions — click to expand"):
            for i,q in enumerate(st.session_state.questions):
                st.markdown(
                    f'<div style="font-family:monospace;font-size:.78rem;padding:.28rem 0;'
                    f'color:#7a92b8;border-bottom:1px solid #0f1420;">'
                    f'<span style="color:#22c55e;margin-right:.5rem;">Q{i+1}</span>{q}</div>',
                    unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    # Clear per-item status so user knows exactly what's missing
    col_st1, col_st2 = st.columns(2)
    with col_st1:
        if has_audio:
            st.markdown('<div class="ok">✅ Step 1 complete — recording uploaded</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="warn">⬆ Step 1 — please upload a recording above</div>', unsafe_allow_html=True)
    with col_st2:
        if has_questions:
            st.markdown(f'<div class="ok">✅ Step 2 complete — {len(st.session_state.questions)} questions loaded</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="warn">⬆ Step 2 — upload a ROPA template or paste questions above</div>', unsafe_allow_html=True)

    ready = has_audio and has_questions

    if st.button("⚡  Transcribe & Analyze ROPA", disabled=not ready, key="btn_ropa"):
        client = get_client()
        prog   = st.progress(0, text="⏳ Transcribing with Groq Whisper-large-v3…")
        af.seek(0)
        try:
            t_data = transcribe_audio(af.read(), af.name, client)
            st.session_state.transcript = t_data.get("text","")
            st.session_state.segments   = t_data.get("segments",[])
            prog.progress(45, text="✅ Transcription done · Analyzing with Llama 3.3 70B…")
        except Exception as e:
            st.error(f"Transcription failed: {e}"); st.stop()

        try:
            result = analyze_ropa(st.session_state.transcript,
                                  st.session_state.segments,
                                  st.session_state.questions, client)
            st.session_state.ropa_result = result
            prog.progress(100, text="✅ Analysis complete!")
            time.sleep(0.4); prog.empty()
        except Exception as e:
            st.error(f"ROPA analysis failed: {e}"); st.stop()

    # ── Results ─────────────────────────────────────────────────────────────
    if st.session_state.ropa_result:
        res   = st.session_state.ropa_result
        ans   = res.get("answers",[])
        compl = res.get("overall_completeness",0)
        high  = sum(1 for a in ans if a.get("confidence")=="HIGH")
        med   = sum(1 for a in ans if a.get("confidence")=="MEDIUM")
        low   = sum(1 for a in ans if a.get("confidence")=="LOW")

        st.markdown(f"""<div class="stat-grid">
          <div class="stat-box"><div class="sn">{len(ans)}</div><div class="sl">Processed</div></div>
          <div class="stat-box"><div class="sn" style="color:#22c55e">{high}</div><div class="sl">High Conf.</div></div>
          <div class="stat-box"><div class="sn" style="color:#f59e0b">{med}</div><div class="sl">Medium</div></div>
          <div class="stat-box"><div class="sn" style="color:#ef4444">{low}</div><div class="sl">Low/Missing</div></div>
          <div class="stat-box"><div class="sn">{compl}%</div><div class="sl">Complete</div></div>
        </div>""", unsafe_allow_html=True)

        st.markdown(f'<div class="ok">✓ {res.get("summary","")}</div>', unsafe_allow_html=True)

        r1,r2 = st.tabs(["📊 ROPA Answers","📝 Transcript"])

        with r1:
            cf = st.selectbox("Filter by confidence",["All","HIGH","MEDIUM","LOW"],key="cf_ropa")
            for a in ans:
                conf = a.get("confidence","LOW")
                if cf!="All" and conf!=cf: continue
                cc = {"HIGH":"c-hi","MEDIUM":"c-md","LOW":"c-lo"}.get(conf,"c-lo")
                bc = {"HIGH":"b-hi","MEDIUM":"b-md","LOW":"b-lo"}.get(conf,"b-lo")
                ts = f'{a.get("timestamp_start","—")} → {a.get("timestamp_end","—")}'
                vq = a.get("verbatim_quote","")
                notes = a.get("notes","")
                vq_html    = f'<div class="verbatim">" {vq} "</div>' if vq else ""
                notes_html = (f'<div style="font-size:.74rem;color:#465370;margin-top:.45rem;'
                              f'border-top:1px solid #0f1420;padding-top:.4rem;">📎 {notes}</div>') if notes else ""
                st.markdown(f"""<div class="qa-card {cc}">
                  <div class="qa-num">Question {a.get('question_index',0)+1}</div>
                  <div class="qa-q">{a.get('question','')}</div>
                  <div class="qa-a">{a.get('answer','')}</div>
                  {vq_html}{notes_html}
                  <div class="qa-meta">
                    <span class="badge {bc}">{conf}</span>
                    <span class="badge b-t">⏱ {ts}</span>
                  </div>
                </div>""", unsafe_allow_html=True)

        with r2:
            if st.session_state.transcript:
                st.text_area("Full Transcript", st.session_state.transcript, height=340, key="tr_v")
                if st.session_state.segments:
                    with st.expander("🕐 Timestamped Segments"):
                        for s in st.session_state.segments:
                            t = f"[{fmt_time(s.get('start'))} → {fmt_time(s.get('end'))}]"
                            st.markdown(
                                f'<div style="font-family:monospace;font-size:.76rem;padding:.22rem 0;'
                                f'color:#5a7090;border-bottom:1px solid #0a0d15;">'
                                f'<span style="color:#22c55e;margin-right:.5rem;">{t}</span>'
                                f'{s.get("text","").strip()}</div>',
                                unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="sh">Export</div>', unsafe_allow_html=True)
        ex1,ex2 = st.columns(2)
        with ex1:
            html_rpt = build_verification_html(res, st.session_state.questions,
                                               st.session_state.audio_meta)
            st.download_button("⬇  Verification Report (HTML)",
                data=html_rpt.encode(),
                file_name=f"ROPA_Verification_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                mime="text/html")
        with ex2:
            rows = [{"Q#":a.get("question_index",0)+1,
                     "Question":a.get("question",""),
                     "Answer":a.get("answer",""),
                     "Confidence":a.get("confidence",""),
                     "Timestamp Start":a.get("timestamp_start",""),
                     "Timestamp End":a.get("timestamp_end",""),
                     "Verbatim Quote":a.get("verbatim_quote",""),
                     "Notes":a.get("notes","")} for a in ans]
            st.download_button("⬇  ROPA Data (CSV)",
                data=pd.DataFrame(rows).to_csv(index=False).encode(),
                file_name=f"ROPA_Data_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — MOM (.docx)
# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown('<div class="card"><div class="card-label">Meeting Metadata</div>'
                '<div class="card-title">Populates the header table in the Word document</div>',
                unsafe_allow_html=True)
    mc1,mc2,mc3 = st.columns(3)
    with mc1:
        proj  = st.text_input("Project / Client Name", placeholder="Privacy Gap Assessment — Acme Corp")
        dept  = st.text_input("Department", placeholder="HR / Finance / IT")
    with mc2:
        fac   = st.text_input("Protiviti / Host Attendee", placeholder="Ms. Smith")
        venue = st.text_input("Venue", value="MS Teams, Virtual Meeting")
    with mc3:
        date_in = st.text_input("Date", placeholder=datetime.now().strftime("%B %d, %Y"))
        time_in = st.text_input("Time", placeholder="10:00 AM")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="card"><div class="card-label">Step 1 · Recording</div>'
                '<div class="card-title">Upload Recording (or reuse from ROPA tab)</div>',
                unsafe_allow_html=True)
    maf = st.file_uploader("MP3 · MP4 · M4A · WAV · OGG · WEBM",
                           type=["mp3","mp4","m4a","wav","ogg","webm","flac","mpeg","mpga"],
                           key="af_mom")
    if maf:
        st.audio(maf)
        st.markdown(f'<div class="ok">✓ {maf.name}</div>', unsafe_allow_html=True)
    reuse = st.session_state.transcript and maf is None
    if reuse:
        st.markdown('<div class="info">ℹ Reusing transcript from ROPA tab — no re-upload needed.</div>',
                    unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    can_run = maf is not None or reuse
    if st.button("⚡  Generate Meeting Minutes (.docx)", disabled=not can_run, key="btn_mom"):
        client = get_client()
        meta   = {"filename": maf.name if maf else "recording",
                  "project_name":proj, "department":dept,
                  "facilitator":fac, "protiviti_attendee":fac,
                  "venue":venue,
                  "date": date_in or datetime.now().strftime("%B %d, %Y"),
                  "time": time_in}

        prog = st.progress(0, text="⏳ Transcribing with Groq Whisper…")
        if maf:
            maf.seek(0)
            try:
                t_data = transcribe_audio(maf.read(), maf.name, client)
                transcript = t_data.get("text","")
                segments   = t_data.get("segments",[])
                if not st.session_state.transcript:
                    st.session_state.transcript = transcript
                    st.session_state.segments   = segments
            except Exception as e: st.error(f"Transcription failed: {e}"); st.stop()
        else:
            transcript = st.session_state.transcript
            segments   = st.session_state.segments

        prog.progress(50, text="✅ Transcript ready · Generating MOM with Llama 3.3 70B…")
        try:
            mom = generate_mom(transcript, segments, meta, client)
            st.session_state.mom_result = mom
            prog.progress(100, text="✅ Done!"); time.sleep(0.4); prog.empty()
        except Exception as e: st.error(f"MOM generation failed: {e}"); st.stop()

    if st.session_state.mom_result:
        mom = st.session_state.mom_result
        st.markdown(f'<div class="ok">✅ "{mom.get("meeting_title","Meeting Minutes")}" generated</div>',
                    unsafe_allow_html=True)

        ov1,ov2 = st.columns(2)
        with ov1:
            st.markdown(f'<div class="mom-box"><div class="mom-lbl">Agenda</div>'
                        f'<div class="mom-val">{mom.get("agenda","—")}</div></div>', unsafe_allow_html=True)
            att = ", ".join(mom.get("attendees_mentioned",[]))
            st.markdown(f'<div class="mom-box"><div class="mom-lbl">Attendees Mentioned</div>'
                        f'<div class="mom-val">{att or "—"}</div></div>', unsafe_allow_html=True)
        with ov2:
            st.markdown(f'<div class="mom-box"><div class="mom-lbl">Next Meeting</div>'
                        f'<div class="mom-val">{mom.get("next_meeting","—")}</div></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="mom-box"><div class="mom-lbl">Next Steps</div>'
                        f'<div class="mom-val">{mom.get("next_steps","—")}</div></div>', unsafe_allow_html=True)

        st.markdown('<div class="sh">Discussion Points</div>', unsafe_allow_html=True)
        for d in mom.get("key_discussion_points",[]):
            ts = (f'<span class="badge b-t" style="margin-left:.4rem;">⏱ {d.get("timestamp","")}</span>'
                  if d.get("timestamp") else "")
            st.markdown(f'<div class="mom-box"><div class="mom-lbl">{d.get("point","")} {ts}</div>'
                        f'<div class="mom-val">{d.get("detail","")}</div></div>', unsafe_allow_html=True)

        st.markdown('<div class="sh">Action Items</div>', unsafe_allow_html=True)
        for item in mom.get("action_items",[]):
            ts = (f'<span class="badge b-t">⏱ {item.get("timestamp","")}</span>'
                  if item.get("timestamp") else "")
            st.markdown(f"""<div class="ai-card">
              <div class="mom-val">{item.get('action','')}</div>
              <div class="qa-meta" style="margin-top:.5rem;">
                <span class="badge" style="background:rgba(124,92,252,.1);color:#7c5cfc;">👤 {item.get('owner','TBD')}</span>
                <span class="badge b-t">📅 {item.get('due_date','TBD')}</span>
                {ts}
              </div>
            </div>""", unsafe_allow_html=True)

        if mom.get("decisions_made"):
            st.markdown('<div class="sh">Decisions Made</div>', unsafe_allow_html=True)
            for d in mom["decisions_made"]:
                st.markdown(f'<div style="padding:.35rem 0;font-size:.87rem;color:#22c55e;'
                            f'border-bottom:1px solid #0a0d15;">✅ {d}</div>', unsafe_allow_html=True)

        if mom.get("questions_raised"):
            st.markdown('<div class="sh">Open Questions</div>', unsafe_allow_html=True)
            for q in mom["questions_raised"]:
                st.markdown(f'<div style="padding:.35rem 0;font-size:.87rem;color:#f59e0b;'
                            f'border-bottom:1px solid #0a0d15;">❓ {q}</div>', unsafe_allow_html=True)

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown('<div class="sh">Download Word Document</div>', unsafe_allow_html=True)

        meta_doc = {"project_name":proj or mom.get("meeting_title",""),
                    "department":dept, "facilitator":fac, "protiviti_attendee":fac,
                    "venue":venue or "MS Teams, Virtual Meeting",
                    "date":date_in or datetime.now().strftime("%B %d, %Y"),
                    "time":time_in, "filename":maf.name if maf else "recording"}

        dl1,dl2 = st.columns(2)
        with dl1:
            try:
                docx_bytes = build_mom_docx(mom, meta_doc)
                st.download_button(
                    "⬇  Download Meeting Minutes (.docx)",
                    data=docx_bytes,
                    file_name=f"MOM_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.markdown('<div class="ok">✓ Word document ready — Protiviti template format</div>',
                            unsafe_allow_html=True)
            except Exception as e:
                st.error(f"DOCX generation failed: {e}")
        with dl2:
            st.download_button(
                "⬇  Download MOM (JSON)",
                data=json.dumps(mom, indent=2).encode(),
                file_name=f"MOM_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json"
            )
