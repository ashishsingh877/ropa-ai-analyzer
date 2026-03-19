"""
mom_docx.py  —  Professional Meeting Minutes Word document generator.
Matches the Protiviti MEETING MINUTES template style exactly.
"""
from __future__ import annotations
import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ──────────────────────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1F, 0x39, 0x64)
C_MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
C_LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK      = RGBColor(0x00, 0x00, 0x00)
C_GRAY       = RGBColor(0x59, 0x59, 0x59)

# ── XML helpers ────────────────────────────────────────────────────────────────
def _shading(cell, hex6: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    old = tcPr.find(qn('w:shd'))
    if old is not None: tcPr.remove(old)
    tcPr.append(shd)

def _row_height(row, twips: int):
    trPr = row._tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(twips))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

def _no_spacing(para):
    pPr = para._p.get_or_add_pPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'),  '0')
    pPr.append(sp)

def _cell_text(cell, text: str, bold=False, italic=False,
               color: RGBColor = C_BLACK, size=10,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    _no_spacing(p)
    r = p.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Arial'

def _set_col_widths(table, widths_twips: list[int]):
    for row in table.rows:
        cells = row.cells
        # de-dup merged cells
        seen = set()
        ci = 0
        for cell in cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                if ci < len(widths_twips):
                    cell.width = Twips(widths_twips[ci])
                ci += 1


# ── Main builder ───────────────────────────────────────────────────────────────
def build_mom_docx(mom_data: dict[str, Any], meta: dict[str, str]) -> bytes:

    doc = Document()

    # Page setup — A4, narrow margins
    for sec in doc.sections:
        sec.page_width    = Cm(21)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # content width in twips: (21 - 1.8*2) cm * 567 twips/cm ≈ 9810
    TW = int((21 - 3.6) * 567)   # ~9810

    # ── TITLE ─────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    _no_spacing(tp)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run('MEETING MINUTES')
    tr.bold = True; tr.font.size = Pt(18)
    tr.font.color.rgb = C_DARK_BLUE; tr.font.name = 'Arial'

    mt = mom_data.get('meeting_title', '')
    if mt:
        sp = doc.add_paragraph()
        _no_spacing(sp)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(mt)
        sr.bold = True; sr.font.size = Pt(12)
        sr.font.color.rgb = C_MID_BLUE; sr.font.name = 'Arial'

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── HELPER: add section-header row spanning all cols ─────────────────────
    def sec_row(tbl, label: str, ncols: int, bg='2E74B5'):
        row = tbl.add_row()
        cell = row.cells[0]
        for i in range(1, ncols):
            cell = cell.merge(row.cells[i])
        _shading(cell, bg)
        _cell_text(cell, label.upper(), bold=True, color=C_WHITE, size=10)
        _row_height(row, 360)

    # ── PARTICULARS TABLE (4 cols) ────────────────────────────────────────────
    pt = doc.add_table(rows=0, cols=4)
    pt.style = 'Table Grid'
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER

    w4 = [int(TW*0.27), int(TW*0.23), int(TW*0.18), int(TW*0.32)]

    # header
    sec_row(pt, 'Meeting Particulars', 4, '1F3964')

    def part_row(l1, v1, l2='', v2=''):
        row = pt.add_row()
        _shading(row.cells[0], 'D6E4F0')
        _cell_text(row.cells[0], l1, bold=True, size=10)
        _cell_text(row.cells[1], v1, size=10)
        if l2:
            _shading(row.cells[2], 'D6E4F0')
            _cell_text(row.cells[2], l2, bold=True, size=10)
            _cell_text(row.cells[3], v2, size=10)
        else:
            m = row.cells[2].merge(row.cells[3])
            _cell_text(m, '', size=10)
        _row_height(row, 320)

    project = meta.get('project_name','')
    dept    = meta.get('department','')
    title_full = f"{project} — {dept}" if dept else project

    part_row('Project Name:',     title_full)
    part_row('Meeting Agenda:',   mom_data.get('agenda',''))
    part_row('Date:',             meta.get('date', datetime.now().strftime('%B %d, %Y')),
             'Time:',             meta.get('time',''))
    part_row('Location / Venue:', meta.get('venue','MS Teams, Virtual Meeting'))
    part_row('Protiviti Attendee:', meta.get('protiviti_attendee', meta.get('facilitator','')))
    part_row('Client Attendee:',  ', '.join(mom_data.get('attendees_mentioned',[])))

    _set_col_widths(pt, w4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── MAIN TABLE (2 cols: # | content) ─────────────────────────────────────
    mt2 = doc.add_table(rows=0, cols=2)
    mt2.style = 'Table Grid'
    mt2.alignment = WD_TABLE_ALIGNMENT.CENTER
    w2 = [int(TW*0.07), int(TW*0.93)]

    def full_row(text, bg='FFFFFF', bold=False, italic=False):
        row = mt2.add_row()
        cell = row.cells[0].merge(row.cells[1])
        _shading(cell, bg)
        _cell_text(cell, text, bold=bold, italic=italic, size=10)
        _row_height(row, 300)

    def num_row(num, text, bg='FFFFFF', bold_text=False):
        row = mt2.add_row()
        _shading(row.cells[0], 'E8F0FA')
        _cell_text(row.cells[0], str(num), bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _shading(row.cells[1], bg)
        _cell_text(row.cells[1], text, size=10, bold=bold_text)
        _row_height(row, 300)

    # ── AGENDA ────────────────────────────────────────────────────────────────
    sec_row(mt2, 'Meeting Agenda', 2)
    full_row(mom_data.get('agenda',''))

    # ── DISCUSSION SUMMARY ───────────────────────────────────────────────────
    sec_row(mt2, 'Discussion Summary', 2)
    dpts = mom_data.get('key_discussion_points', [])
    if dpts:
        for i, d in enumerate(dpts, 1):
            ts  = f"  [{d.get('timestamp','')}]" if d.get('timestamp') else ''
            num_row(i, f"{d.get('point','')}{ts}", bg='EDF3FB', bold_text=True)
            if d.get('detail'):
                full_row(d['detail'], bg='FFFFFF')
    else:
        full_row('No discussion points recorded.', italic=True)

    # ── DECISIONS ────────────────────────────────────────────────────────────
    decs = mom_data.get('decisions_made', [])
    if decs:
        sec_row(mt2, 'Decisions Made', 2)
        for i, d in enumerate(decs, 1):
            num_row(i, d, bg='E2EFDA')

    # ── ACTION ITEMS ─────────────────────────────────────────────────────────
    ais = mom_data.get('action_items', [])
    if ais:
        sec_row(mt2, 'Action Items', 2)
        # sub-header
        sh = mt2.add_row()
        _shading(sh.cells[0], '1F3964')
        _cell_text(sh.cells[0], '#', bold=True, color=C_WHITE, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _shading(sh.cells[1], '1F3964')
        p = sh.cells[1].paragraphs[0]; p.clear(); _no_spacing(p)
        for lbl in ['Action', '        Owner', '        Due Date']:
            rr = p.add_run(lbl); rr.bold = True
            rr.font.size = Pt(9); rr.font.color.rgb = C_WHITE; rr.font.name = 'Arial'
        _row_height(sh, 280)

        for i, item in enumerate(ais, 1):
            bg = 'FFF9E6' if i % 2 else 'FFF0CC'
            row = mt2.add_row()
            _shading(row.cells[0], bg)
            _cell_text(row.cells[0], str(i), bold=True, size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _shading(row.cells[1], bg)

            cp = row.cells[1].paragraphs[0]; cp.clear(); _no_spacing(cp)
            def _r(txt, bold=False, color=C_BLACK):
                rr = cp.add_run(str(txt))
                rr.bold = bold; rr.font.size = Pt(10)
                rr.font.color.rgb = color; rr.font.name = 'Arial'

            _r(item.get('action',''))
            _r('\n  Owner: ', True, C_MID_BLUE)
            _r(item.get('owner','TBD'))
            _r('   |   Due: ', True, C_MID_BLUE)
            _r(item.get('due_date','TBD'))
            ts2 = item.get('timestamp','')
            if ts2:
                _r(f'   [{ts2}]', color=RGBColor(0x80,0x80,0x80))
            _row_height(row, 420)

    # ── OPEN QUESTIONS ───────────────────────────────────────────────────────
    qrs = mom_data.get('questions_raised', [])
    if qrs:
        sec_row(mt2, 'Open Questions / Points to Clarify', 2)
        for i, q in enumerate(qrs, 1):
            num_row(i, q, bg='FFF9E6')

    # ── NEXT STEPS ───────────────────────────────────────────────────────────
    sec_row(mt2, 'Next Steps', 2)
    full_row(mom_data.get('next_steps','To be confirmed.'))

    nm = mom_data.get('next_meeting')
    if nm:
        sec_row(mt2, 'Next Meeting', 2)
        full_row(nm)

    _set_col_widths(mt2, w2)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    _no_spacing(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f'Confidential  ·  Generated by ROPA AI Analyzer  ·  '
        f'{datetime.now().strftime("%d %B %Y, %H:%M")}'
    )
    fr.font.size = Pt(8); fr.italic = True
    fr.font.color.rgb = RGBColor(0xAA,0xAA,0xAA); fr.font.name = 'Arial'

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
